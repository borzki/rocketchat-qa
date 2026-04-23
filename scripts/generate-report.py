"""Generate midterm-report.docx from structured data."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ── Title Page ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Midterm Report\nQA Implementation & Empirical Analysis')
run.font.size = Pt(26)
run.bold = True

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Rocket.Chat — Open-Source Team Communications Platform')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Course: ').bold = False
info.add_run('Advanced QA\n').bold = False
info.add_run('Date: ').bold = False
info.add_run('2026-04-11\n').bold = False
info.add_run('Repository: ').bold = False
info.add_run('github.com/illus1um/rocketchat-qa\n').bold = False

doc.add_page_break()

# ── 1. System Description ──
doc.add_heading('1. System Description', level=1)

doc.add_heading('1.1 Architecture', level=2)
doc.add_paragraph(
    'Rocket.Chat is an open-source team communications platform (self-hosted alternative to Slack). '
    'It follows a monolithic architecture built on the Meteor full-stack reactive framework.'
)

add_table(
    ['Layer', 'Technology'],
    [
        ['Language', 'TypeScript / JavaScript'],
        ['Framework', 'Meteor (full-stack reactive)'],
        ['Runtime', 'Node.js'],
        ['Database', 'MongoDB (single-node replica set)'],
        ['Real-time Layer', 'WebSocket / DDP (Distributed Data Protocol)'],
        ['Deployment', 'Docker Compose'],
        ['API', 'REST API (150+ endpoints)'],
    ]
)
doc.add_paragraph()

doc.add_heading('1.2 Key Functionalities', level=2)
for item in [
    'Real-time messaging across channels, groups, and direct messages',
    'REST API for integrations, bots, and automation',
    'User authentication (local credentials, OAuth, LDAP, SAML, 2FA)',
    'Role-based access control (RBAC)',
    'File upload and sharing',
    'End-to-end encryption',
    'Omnichannel / Livechat for customer support',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.3 Deployment Under Test', level=2)
add_table(
    ['Parameter', 'Value'],
    [
        ['Edition', 'Community (open-source, self-hosted)'],
        ['Deployment', 'Docker Compose (docker-compose.yml)'],
        ['Services', 'Rocket.Chat (latest) + MongoDB 8 (replica set rs0)'],
        ['Base URL', 'http://localhost:3000'],
        ['Admin setup', 'Automated via environment variables'],
    ]
)
doc.add_paragraph()

# ── 2. Methodology ──
doc.add_heading('2. Methodology', level=1)

doc.add_heading('2.1 Risk-Based Testing Approach', level=2)
doc.add_paragraph(
    'Testing was prioritized using a quantitative risk model from Assignment 1: '
    'Risk Score = Probability (P) x Impact (I), both rated 1-5. '
    'Modules scoring 12+ were prioritized for automation.'
)

doc.add_paragraph('Top 3 high-risk modules targeted for midterm testing:')
add_table(
    ['Module', 'Risk Score', 'Priority'],
    [
        ['Real-time Messaging', '20', 'P0 — Critical'],
        ['REST API', '16', 'P0 — Critical'],
        ['Authentication & Authorization', '15', 'P0 — Critical'],
    ]
)
doc.add_paragraph()

doc.add_heading('2.2 Test Design Strategy', level=2)
doc.add_paragraph('Tests were designed to cover four mandatory categories:')
for cat, desc in [
    ('Failure Scenarios', 'how the system behaves when things go wrong (invalid credentials, expired tokens, error codes 400/401)'),
    ('Edge Cases', 'boundary conditions (empty input, 50KB payloads, special characters, injection-like input)'),
    ('Concurrency / Race Conditions', 'simultaneous requests, parallel API calls, rapid sequential actions'),
    ('Invalid User Behavior', 'unauthorized access to restricted endpoints, skipping required steps'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{cat} — ').bold = True
    p.add_run(desc)

doc.add_heading('2.3 Automation Tools', level=2)
add_table(
    ['Tool', 'Purpose', 'Tests'],
    [
        ['Jest', 'Test runner and assertions', '37'],
        ['Axios', 'HTTP client for API calls', '—'],
        ['Docker Compose', 'Environment orchestration', '—'],
        ['GitHub Actions', 'CI/CD pipeline', '—'],
    ]
)
doc.add_paragraph()

# ── 3. Automation Implementation ──
doc.add_heading('3. Automation Implementation', level=1)

doc.add_heading('3.1 Test Structure', level=2)
add_table(
    ['File', 'Tests', 'Description'],
    [
        ['auth.test.js', '5', 'Basic authentication CRUD'],
        ['auth-edge.test.js', '8 (NEW)', 'Failure, edge, injection, RBAC'],
        ['channels.test.js', '4', 'Basic channel CRUD'],
        ['channels-edge.test.js', '7 (NEW)', 'Edge cases, duplicates, concurrency'],
        ['messaging.test.js', '4', 'Basic messaging CRUD'],
        ['messaging-edge.test.js', '9 (NEW)', 'Edge, failure, concurrency'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('Total: 37 tests across 6 test suites. 24 new tests added for midterm.')

doc.add_heading('3.2 Test Case Inventory — Authentication (13 tests)', level=2)
add_table(
    ['Test ID', 'Type', 'Description'],
    [
        ['AUTH-01', 'Happy path', 'Login with valid credentials'],
        ['AUTH-02', 'Failure', 'Reject invalid password'],
        ['AUTH-03', 'Failure', 'Reject non-existent user'],
        ['AUTH-04', 'Happy path', 'GET /me with valid token'],
        ['AUTH-05', 'Happy path', 'Logout invalidates session'],
        ['TC-AUTH-FAIL-01', 'Failure', 'Reject expired/invalid token'],
        ['TC-AUTH-FAIL-02', 'Failure', 'Reject missing auth headers'],
        ['TC-AUTH-FAIL-03', 'Edge case', 'Reject empty password'],
        ['TC-AUTH-FAIL-04', 'Edge case', 'Reject empty username'],
        ['TC-AUTH-EDGE-01', 'Edge case', 'Reject NoSQL injection input'],
        ['TC-AUTH-EDGE-02', 'Edge case', 'Reject 10000-char username'],
        ['TC-AUTH-EDGE-03', 'Edge case', 'XSS input not reflected in response'],
        ['TC-AUTH-INVALID-01', 'Invalid behavior', 'Regular user cannot access admin endpoints'],
    ]
)
doc.add_paragraph()

doc.add_heading('3.3 Test Case Inventory — Channels (11 tests)', level=2)
add_table(
    ['Test ID', 'Type', 'Description'],
    [
        ['CHAN-01', 'Happy path', 'Create public channel'],
        ['CHAN-02', 'Happy path', 'List channels'],
        ['CHAN-03', 'Happy path', 'Get channel info'],
        ['CHAN-04', 'Happy path', 'Delete channel'],
        ['TC-CHAN-EDGE-01', 'Edge case', 'Reject special chars in channel name'],
        ['TC-CHAN-EDGE-02', 'Edge case', 'Reject empty channel name'],
        ['TC-CHAN-EDGE-03', 'Edge case', 'Reject duplicate channel creation'],
        ['TC-CHAN-FAIL-01', 'Failure', 'Reject unauthenticated creation'],
        ['TC-CHAN-FAIL-02', 'Failure', 'Error for non-existent channel info'],
        ['TC-CHAN-FAIL-03', 'Failure', 'Reject deleting already deleted channel'],
        ['TC-CHAN-CONC-01', 'Concurrency', '5 simultaneous channel creations'],
    ]
)
doc.add_paragraph()

doc.add_heading('3.4 Test Case Inventory — Messaging (13 tests)', level=2)
add_table(
    ['Test ID', 'Type', 'Description'],
    [
        ['MSG-01', 'Happy path', 'Send message to channel'],
        ['MSG-02', 'Happy path', 'Retrieve message history'],
        ['MSG-03', 'Happy path', 'Edit existing message'],
        ['MSG-04', 'Happy path', 'Delete message'],
        ['TC-MSG-EDGE-01', 'Edge case', 'Handle empty message'],
        ['TC-MSG-EDGE-02', 'Edge case', 'Special characters & unicode'],
        ['TC-MSG-EDGE-03', 'Edge case', 'Very large message (50KB)'],
        ['TC-MSG-FAIL-01', 'Failure', 'Reject message to non-existent room'],
        ['TC-MSG-FAIL-02', 'Failure', 'Reject unauthenticated send'],
        ['TC-MSG-FAIL-03', 'Failure', 'Reject updating non-existent message'],
        ['TC-MSG-FAIL-04', 'Failure', 'Reject deleting non-existent message'],
        ['TC-MSG-CONC-01', 'Concurrency', '10 simultaneous messages — no data loss'],
        ['TC-MSG-CONC-02', 'Concurrency', 'Rapid send-then-delete consistency'],
    ]
)
doc.add_paragraph()

doc.add_heading('3.5 CI/CD Setup', level=2)
doc.add_paragraph(
    'Pipeline: GitHub Actions (.github/workflows/ci.yml)\n'
    'Trigger: Every push to main, pull requests, manual dispatch'
)
doc.add_paragraph('Pipeline flow:')
for step in [
    'Push to GitHub triggers workflow',
    'Checkout repository, setup Node.js 20, install dependencies',
    'Docker Compose starts Rocket.Chat + MongoDB',
    'Wait for MongoDB health check (replica set auto-init)',
    'Wait for Rocket.Chat API to respond (up to 360 seconds)',
    'Verify admin login via REST API',
    'Run 37 API tests with coverage reporting',
    'Evaluate Quality Gates (3 automated gates)',
    'Upload coverage artifacts',
    'Docker Compose down (cleanup)',
]:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('3.6 Quality Gates Definition', level=2)
add_table(
    ['Gate', 'Threshold', 'Rationale'],
    [
        ['Gate 1: Test Pass Rate', '>= 90%', 'Ensures overall test health'],
        ['Gate 2: Zero Critical Failures', '0 failures', 'Security tests must never fail'],
        ['Gate 3: All Test Suites Pass', '0 failed suites', 'No entire module can be broken'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('The pipeline fails automatically if any gate is not met.')

# ── 4. Results ──
doc.add_heading('4. Results', level=1)

doc.add_heading('4.1 Test Execution Summary', level=2)
add_table(
    ['Metric', 'Value'],
    [
        ['Total tests', '37'],
        ['Passed', '37'],
        ['Failed', '0'],
        ['Test suites', '6'],
        ['Execution time (local)', '~2.5 seconds'],
        ['Execution time (CI)', '~1 min 32 sec (incl. Docker startup)'],
        ['CI Run', 'Run #9 — Success'],
    ]
)
doc.add_paragraph()

doc.add_heading('4.2 Quality Gate Results', level=2)
add_table(
    ['Gate', 'Result', 'Value'],
    [
        ['Gate 1 — Pass Rate >= 90%', 'PASSED', '100%'],
        ['Gate 2 — Zero critical failures', 'PASSED', '0 failures'],
        ['Gate 3 — All test suites pass', 'PASSED', '6/6 suites'],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('ALL QUALITY GATES PASSED')
run.bold = True
run.font.size = Pt(12)

doc.add_heading('4.3 Results by Module', level=2)
add_table(
    ['Module', 'Tests', 'Passed', 'Failed', 'Pass Rate'],
    [
        ['Authentication & Authorization', '13', '13', '0', '100%'],
        ['REST API — Channels', '11', '11', '0', '100%'],
        ['REST API — Messaging', '13', '13', '0', '100%'],
        ['Total', '37', '37', '0', '100%'],
    ]
)
doc.add_paragraph()

doc.add_heading('4.4 Results by Category', level=2)
add_table(
    ['Category', 'Tests', 'Passed', 'Failed'],
    [
        ['Happy path (CRUD)', '13', '13', '0'],
        ['Failure scenarios', '10', '10', '0'],
        ['Edge cases', '9', '9', '0'],
        ['Concurrency', '3', '3', '0'],
        ['Invalid user behavior', '2', '2', '0'],
    ]
)
doc.add_paragraph()

doc.add_heading('4.5 Defect Detection', level=2)
add_table(
    ['#', 'Test', 'Finding', 'Severity', 'Module'],
    [
        ['1', 'TC-MSG-EDGE-01', 'RC accepts empty messages without error', 'Low', 'Messaging'],
        ['2', 'TC-MSG-EDGE-03', '50KB message accepted — no server-side size limit', 'Medium', 'Messaging / DB'],
        ['3', 'TC-AUTH-INVALID-01', 'Open user registration via API without explicit policy', 'Low', 'Authentication'],
    ]
)
doc.add_paragraph()

doc.add_paragraph('Defect mapping to risk levels:')
add_table(
    ['Defect', 'Risk Module', 'Risk Score', 'Impact'],
    [
        ['Empty messages accepted', 'Real-time Messaging', '20 (P0)', 'Low — cosmetic, affects P0 module'],
        ['No message size limit', 'REST API + Database', '16/15 (P0/P1)', 'Medium — DoS potential'],
        ['Open user registration', 'Authentication', '15 (P0)', 'Low — default config concern'],
    ]
)
doc.add_paragraph()

doc.add_heading('4.6 API Endpoint Coverage', level=2)
doc.add_paragraph(
    'Since tests are integration tests hitting a running Rocket.Chat instance via HTTP (black-box), '
    'traditional code coverage (Istanbul/JaCoCo) does not apply. '
    'API endpoint coverage is the closest meaningful metric:'
)
add_table(
    ['API Area', 'Total Endpoints', 'Tested', 'Coverage'],
    [
        ['Authentication (/login, /logout, /me)', '3', '3', '100%'],
        ['Channels (create, list, info, delete)', '4', '4', '100%'],
        ['Messaging (send, history, update, delete)', '4', '4', '100%'],
        ['User management (register, list, delete)', '3', '2', '67%'],
        ['High-risk modules total', '14', '13', '93%'],
    ]
)
doc.add_paragraph()

doc.add_heading('4.7 Flaky Test Analysis', level=2)
add_table(
    ['Metric', 'Value'],
    [
        ['CI runs analyzed', '2 (runs #8 and #9)'],
        ['Flaky tests detected', '0'],
        ['Flaky test rate', '0%'],
    ]
)
doc.add_paragraph()

doc.add_heading('4.8 Efficiency Metrics', level=2)
add_table(
    ['Metric', 'Assignment 1', 'Midterm', 'Change'],
    [
        ['Total tests', '13', '37', '+185%'],
        ['Test execution time', '~1.1 sec', '~2.5 sec', '+127%'],
        ['Tests per second', '11.8', '14.8', '+25% improvement'],
        ['CI pipeline time', 'N/A', '92 sec', 'New'],
        ['Defects found', '0', '3', '+3'],
    ]
)
doc.add_paragraph()

# ── 5. Refined Risk-Based Testing Strategy (Task 1) ──
doc.add_heading('5. Refined Risk-Based Testing Strategy', level=1)

doc.add_heading('5.1 Re-evaluation of High-Risk Components', level=2)
add_table(
    ['Module', 'Original Score', 'Observed Issues', 'Updated Score', 'Justification'],
    [
        [
            'Real-time Messaging',
            'P=4, I=5\nScore=20',
            'Empty messages accepted; 50KB stored without limit; 10 concurrent messages OK',
            'P=4, I=5\nScore=20',
            'Risk remains Critical. Lack of input validation confirms high defect probability.',
        ],
        [
            'REST API',
            'P=4, I=4\nScore=16',
            'All CRUD stable; consistent error codes; 5 concurrent channel creations OK',
            'P=3, I=4\nScore=12',
            'Probability reduced 4->3. API demonstrated consistent behavior. Moved P0 to P1.',
        ],
        [
            'Authentication',
            'P=3, I=5\nScore=15',
            'Injection rejected; XSS not reflected; expired tokens rejected; RBAC partial',
            'P=3, I=5\nScore=15',
            'Unchanged. Security tests passed but untested 2FA and open registration keep P at 3.',
        ],
    ]
)
doc.add_paragraph()

doc.add_heading('5.2 Evidence from Automation Runs', level=2)

p = doc.add_paragraph()
p.add_run('A. Failed Test Cases: ').bold = True
p.add_run('No test failures observed. All 37 tests passed consistently.')

p = doc.add_paragraph()
p.add_run('B. Flaky Tests: ').bold = True
p.add_run('No flaky tests detected. 0% flaky rate across all executions.')

p = doc.add_paragraph()
p.add_run('C. Coverage Gaps:').bold = True
doc.add_paragraph('File upload endpoints — not tested (P1 module, deferred)', style='List Bullet')
doc.add_paragraph('E2E encryption API — not tested (requires browser-side crypto)', style='List Bullet')
doc.add_paragraph('Omnichannel/Livechat — not tested (P2 module, out of scope)', style='List Bullet')
doc.add_paragraph('2FA, OAuth, LDAP authentication flows — not tested', style='List Bullet')

p = doc.add_paragraph()
p.add_run('D. Unexpected System Behavior:').bold = True
doc.add_paragraph('Empty messages accepted — POST /chat.sendMessage with msg:"" returns 200 OK', style='List Bullet')
doc.add_paragraph('No message size limit — 50,000-character message accepted and stored in full', style='List Bullet')
doc.add_paragraph('Open user registration — POST /users.register available without authentication', style='List Bullet')

doc.add_heading('5.3 Risk Dimension Mapping', level=2)
add_table(
    ['Module', 'Likelihood', 'Impact', 'Detectability'],
    [
        ['Real-time Messaging', 'Unchanged (edge cases confirmed)', 'Critical (core feature)', 'Improved (13 tests)'],
        ['REST API', 'Reduced (consistent error handling)', 'Major (integrations depend on it)', 'Improved (11 tests)'],
        ['Authentication', 'Unchanged (injection safe, gaps remain)', 'Critical (security module)', 'Improved (13 tests)'],
    ]
)
doc.add_paragraph()

# ── 6. Comparative Analysis (Task 4) ──
doc.add_heading('6. Comparative Analysis', level=1)

doc.add_heading('6.1 Planned vs Actual', level=2)
add_table(
    ['Aspect', 'Planned (A1)', 'Actual (Midterm)', 'Gap'],
    [
        ['API tests', '13 (auth, channels, messaging)', '37 (+24 edge/failure/concurrency)', 'Exceeded by 185%'],
        ['UI E2E tests', '5 smoke tests (Playwright)', '0 (focused on API depth)', 'Deferred — API provides more value'],
        ['Newman/Postman', '7 requests, 14 assertions', 'Not included', 'Redundant — Jest covers same endpoints'],
        ['Performance', 'JMeter planned', 'Not executed', 'Concurrency tests partially cover gap'],
        ['CI/CD pipeline', 'Not operational', 'Fully operational with quality gates', 'Major improvement'],
        ['Quality gates', 'Not defined', '3 gates automated', 'New deliverable'],
        ['Defects found', '0', '3', 'Improved detection via edge cases'],
    ]
)
doc.add_paragraph()

doc.add_heading('6.2 Key Insights', level=2)

p = doc.add_paragraph()
p.add_run('Incorrect Assumptions in Planning:').bold = True

doc.add_paragraph(
    'UI E2E was overvalued in A1. Rocket.Chat UI is volatile across versions. '
    'API-level testing proved far more stable and provided better coverage of high-risk modules.',
    style='List Bullet'
)
doc.add_paragraph(
    'Newman was redundant. Jest allows programmatic setup/teardown, concurrent assertions, '
    'and JSON output for quality gates — Newman does not.',
    style='List Bullet'
)
doc.add_paragraph(
    'Performance testing required dedicated infrastructure. Local Docker and CI runners '
    'are insufficient for meaningful benchmarks.',
    style='List Bullet'
)

p = doc.add_paragraph()
p.add_run('Missing Test Scenarios (Identified During Midterm):').bold = True
doc.add_paragraph('Token expiration/rotation — tested invalid tokens but not naturally expired ones', style='List Bullet')
doc.add_paragraph('Rate limiting — no tests verify API enforces rate limits', style='List Bullet')
doc.add_paragraph('File upload security — P1 module remains completely untested', style='List Bullet')
doc.add_paragraph('WebSocket/DDP real-time protocol — P0 module tested only through REST API', style='List Bullet')

# ── 7. Discussion ──
doc.add_heading('7. Discussion', level=1)

doc.add_heading('7.1 What Worked', level=2)
doc.add_paragraph(
    'Risk-based test prioritization was effective. By focusing on P0 modules, we achieved '
    '93% endpoint coverage on high-risk areas with 37 tests.',
    style='List Number'
)
doc.add_paragraph(
    'API integration testing is highly stable. All 37 tests pass deterministically with 0% flaky rate. '
    'Unlike UI tests which depend on CSS selectors and rendering timing, API tests hit stable REST endpoints.',
    style='List Number'
)
doc.add_paragraph(
    'Docker Compose auto-setup eliminated manual intervention. Environment variables skip the Setup Wizard '
    'and create the admin user automatically — both locally and in CI.',
    style='List Number'
)
doc.add_paragraph(
    'Quality gates provide objective go/no-go criteria. The three automated gates turn the CI pipeline '
    'into a decision-making tool rather than just a test runner.',
    style='List Number'
)

doc.add_heading('7.2 What Didn\'t Work', level=2)
doc.add_paragraph(
    'Traditional code coverage is not applicable. Since we test a running instance via HTTP (black-box), '
    'we cannot measure line/branch coverage. API endpoint coverage is the closest meaningful metric.',
    style='List Number'
)
doc.add_paragraph(
    'Performance testing gap persists. JMeter was planned in A1 but never executed. '
    'Concurrency tests provide some load validation, but proper benchmarking requires dedicated infrastructure.',
    style='List Number'
)

doc.add_heading('7.3 Unexpected Findings', level=2)
doc.add_paragraph(
    'Rocket.Chat accepts empty messages — minor UI/UX defect not predicted in risk assessment.',
    style='List Number'
)
doc.add_paragraph(
    'No server-side message size limit — 50KB message accepted without error. '
    'Validates P0 classification of REST API and Database modules.',
    style='List Number'
)
doc.add_paragraph(
    'MongoDB health check can self-initialize replica sets — eliminates need for separate init container, '
    'making Docker setup more robust and CI-friendly.',
    style='List Number'
)

doc.add_heading('7.4 Improvements for Next Phase', level=2)
doc.add_paragraph('Add WebSocket/DDP tests for Real-time Messaging module (P0, score 20)', style='List Bullet')
doc.add_paragraph('Add file upload security tests for File Upload module (P1, score 12)', style='List Bullet')
doc.add_paragraph('Implement rate limiting tests to verify API abuse protection', style='List Bullet')
doc.add_paragraph('Add performance benchmarks using k6 or Artillery (lighter than JMeter, CI-friendly)', style='List Bullet')
doc.add_paragraph('Test 2FA enforcement and OAuth flows for Authentication module', style='List Bullet')

# ── Save ──
output_path = 'docs/midterm-report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
